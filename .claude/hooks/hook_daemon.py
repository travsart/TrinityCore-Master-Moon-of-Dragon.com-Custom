#!/usr/bin/env python3
"""VoxCore Hook Daemon — persistent asyncio HTTP server for Claude Code hooks.

Eliminates the ~100ms Python cold start that every command-type hook pays.
Listens on 127.0.0.1:19484 for POSTs from Claude Code's type:"http" hooks.

Stdlib only. Python 3.12+. Windows tested.

Architecture:
    Claude Code 2.1.97
    └── type: "http" hook
        └── POST http://127.0.0.1:19484/hook/<name>
            └── ROUTE_TABLE[path] → handler(data) → {} or {"decision": "block", ...}

Each handler is a single async function. Sync blocking handlers return their
result immediately. Async advisory handlers spawn an asyncio.create_task() and
return {} so Claude Code is never blocked on background work.

Plan: C:/Users/atayl/.claude/plans/imperative-drifting-hoare.md
"""

import asyncio
import ctypes
import http.client
import json
import logging
import os
import re
import signal
import sqlite3
import subprocess
import sys
import time
import unicodedata
from collections import Counter
from datetime import date, datetime, timedelta, timezone
from logging.handlers import RotatingFileHandler
from pathlib import Path

# ─── SECTION A — Constants ──────────────────────────────────────────────

PORT = 19484
HOST = "127.0.0.1"

PROJECT_DIR = Path(__file__).resolve().parent.parent.parent  # .claude/hooks/.. → VoxCore root
HOOKS_DIR = PROJECT_DIR / ".claude" / "hooks"
GATE_STATUS_FILE = PROJECT_DIR / ".claude" / "release-gate-status.json"
DEADLINES_FILE = PROJECT_DIR / ".claude" / "deadlines.json"

USER_CLAUDE = Path.home() / ".claude"
PID_FILE = USER_CLAUDE / "hook_daemon.pid"
LOG_FILE = USER_CLAUDE / "hook_daemon.log"
STATS_FILE = USER_CLAUDE / "session-stats.jsonl"
PRECOMPACT_STATE = USER_CLAUDE / "precompact-state.json"

DAEMON_VERSION = "1.1.0"
START_TIME = time.time()
IN_FLIGHT: "set[asyncio.Task]" = set()
SHUTDOWN_EVENT: "asyncio.Event | None" = None  # set in main()

# ── Feature #1: Codebase DB auto-refresh dirty flag ─────────────────────
CODEBASE_DIRTY_FLAG = USER_CLAUDE / "codebase_dirty.flag"
CODEBASE_DB_PATH = PROJECT_DIR / ".cache" / "codebase.db"

# ── Feature #2: Digest staleness tracking ────────────────────────────────
DIGEST_CACHE_DIR = PROJECT_DIR / ".cache" / "source_digests"

# ── Feature #9: Triad reminder — track new C++ lines per session ────────
SESSION_CPP_NEW_LINES = 0

# ─── SECTION B — Logging ────────────────────────────────────────────────

USER_CLAUDE.mkdir(parents=True, exist_ok=True)
log = logging.getLogger("hook_daemon")
log.setLevel(logging.INFO)
_h = RotatingFileHandler(
    LOG_FILE, maxBytes=5_000_000, backupCount=3, encoding="utf-8"
)
_h.setFormatter(logging.Formatter("%(asctime)s %(levelname)-7s %(message)s"))
log.addHandler(_h)

# ─── SECTION C — HTTP server core ───────────────────────────────────────

async def parse_http_request(reader: asyncio.StreamReader) -> "tuple[str, str, bytes]":
    """Parse an HTTP/1.1 request. Returns (method, path, body)."""
    line = await reader.readline()
    if not line:
        raise ConnectionError("empty request")
    parts = line.decode("ascii", "replace").rstrip("\r\n").split(" ", 2)
    if len(parts) < 3:
        raise ValueError(f"malformed request line: {line!r}")
    method, path, _version = parts
    headers: "dict[str, str]" = {}
    while True:
        hline = await reader.readline()
        if hline in (b"\r\n", b"\n", b""):
            break
        name, _, value = hline.decode("ascii", "replace").partition(":")
        headers[name.strip().lower()] = value.strip()
    body = b""
    content_length = int(headers.get("content-length", "0") or 0)
    if content_length > 0:
        body = await reader.readexactly(content_length)
    return method, path, body


async def send_response(
    writer: asyncio.StreamWriter, status: int, body: "dict | None"
) -> None:
    """Send an HTTP/1.1 response with a JSON body."""
    body_bytes = json.dumps(body if body is not None else {}, ensure_ascii=False).encode(
        "utf-8"
    )
    status_text = {
        200: "OK",
        400: "Bad Request",
        404: "Not Found",
        500: "Internal Server Error",
    }.get(status, "OK")
    response = (
        f"HTTP/1.1 {status} {status_text}\r\n"
        f"Content-Type: application/json; charset=utf-8\r\n"
        f"Content-Length: {len(body_bytes)}\r\n"
        f"Connection: close\r\n\r\n"
    ).encode("ascii") + body_bytes
    writer.write(response)
    await writer.drain()


async def client_connected(
    reader: asyncio.StreamReader, writer: asyncio.StreamWriter
) -> None:
    """Handle one HTTP request."""
    try:
        method, path, body = await asyncio.wait_for(
            parse_http_request(reader), timeout=10.0
        )
        # Health endpoint
        if method == "GET" and path == "/health":
            uptime = time.time() - START_TIME
            await send_response(
                writer,
                200,
                {
                    "status": "ok",
                    "pid": os.getpid(),
                    "uptime": round(uptime, 2),
                    "version": DAEMON_VERSION,
                    "in_flight": len(IN_FLIGHT),
                },
            )
            return
        # Shutdown endpoint (POST /shutdown) — Windows-friendly graceful stop
        # since taskkill cannot send SIGTERM to detached Python processes.
        if method == "POST" and path == "/shutdown":
            await send_response(writer, 200, {"status": "shutting down"})
            log.info("Shutdown requested via HTTP")
            if SHUTDOWN_EVENT is not None:
                SHUTDOWN_EVENT.set()
            return
        # Route lookup (POST only for hooks)
        if method != "POST":
            await send_response(writer, 400, {"error": f"method not allowed: {method}"})
            return
        handler = ROUTE_TABLE.get(path)
        if handler is None:
            log.warning("unknown path: %s", path)
            await send_response(writer, 404, {"error": f"unknown path: {path}"})
            return
        # Parse body
        try:
            data = json.loads(body) if body else {}
        except json.JSONDecodeError as e:
            log.warning("invalid JSON body on %s: %s", path, e)
            await send_response(writer, 200, {})  # graceful pass-through
            return
        # Run handler with error isolation
        task = asyncio.current_task()
        if task is not None:
            IN_FLIGHT.add(task)
        try:
            result = await handler(data)
            await send_response(writer, 200, result if result else {})
        except Exception:
            log.exception("handler error on %s", path)
            await send_response(writer, 200, {})  # pass-through on error
        finally:
            if task is not None:
                IN_FLIGHT.discard(task)
    except asyncio.TimeoutError:
        log.warning("request parse timeout")
    except ConnectionError as e:
        log.debug("connection error: %s", e)
    except Exception:
        log.exception("connection handler error")
    finally:
        try:
            writer.close()
            await writer.wait_closed()
        except Exception:
            pass


# ─── SECTION D — Helpers (shared) ───────────────────────────────────────

def _get_tool_input(data: dict) -> dict:
    """Safe accessor: tool_input must be a dict."""
    ti = data.get("tool_input")
    return ti if isinstance(ti, dict) else {}


def _normalize_path(p: str) -> str:
    return p.replace("\\", "/").lower()


def _read_jsonl_recent(path: Path, minutes: int) -> "list[dict]":
    """Read JSONL entries newer than `minutes` ago. Empty list on any error."""
    out: "list[dict]" = []
    cutoff = datetime.now(timezone.utc) - timedelta(minutes=minutes)
    try:
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    entry = json.loads(line)
                except json.JSONDecodeError:
                    continue
                ts_str = entry.get("timestamp", "")
                try:
                    ts = datetime.fromisoformat(ts_str)
                    if ts < cutoff:
                        continue
                except (ValueError, TypeError):
                    continue
                out.append(entry)
    except FileNotFoundError:
        pass
    except Exception:
        log.exception("_read_jsonl_recent error")
    return out


def _spawn_detached(args: "list[str]", log_to: "Path | None" = None) -> None:
    """Spawn a detached subprocess on Windows. Fire-and-forget."""
    DETACHED_PROCESS = 0x00000008
    CREATE_NEW_PROCESS_GROUP = 0x00000200
    try:
        if log_to is not None:
            log_to.parent.mkdir(parents=True, exist_ok=True)
            f = open(log_to, "a", encoding="utf-8")
            stdout = f
            stderr = f
        else:
            stdout = subprocess.DEVNULL
            stderr = subprocess.DEVNULL
        subprocess.Popen(
            args,
            stdout=stdout,
            stderr=stderr,
            creationflags=DETACHED_PROCESS | CREATE_NEW_PROCESS_GROUP,
            close_fds=True,
        )
    except Exception:
        log.exception("_spawn_detached failed: %s", args)


# ═══════════════════════════════════════════════════════════════════════
# SECTION E — SYNC BLOCKING HANDLERS
# Each returns a dict matching CC's hook output schema:
#   blocking: {"decision": "block", "reason": "..."}
#   context:  {"additionalContext": "..."}
#   no-op:    {}
# ═══════════════════════════════════════════════════════════════════════

# ── handle_sql_safety ────────────────────────────────────────────────────
DANGEROUS_SQL_PATTERNS = [
    (re.compile(r"\bDROP\s+TABLE\b", re.I), "DROP TABLE"),
    (re.compile(r"\bDROP\s+DATABASE\b", re.I), "DROP DATABASE"),
    (re.compile(r"\bTRUNCATE\b", re.I), "TRUNCATE"),
    (re.compile(r"\bDELETE\s+FROM\s+\S+\s*;", re.I), "DELETE without WHERE clause"),
    (re.compile(r"\bDELETE\s+FROM\s+\S+\s*$", re.I), "DELETE without WHERE clause"),
    (re.compile(r"\bALTER\s+TABLE\s+\S+\s+DROP\s+COLUMN\b", re.I), "ALTER TABLE DROP COLUMN"),
]
SQL_SAFE_OVERRIDES = [
    re.compile(r"DROP\s+TABLE\s+IF\s+EXISTS.*CREATE\s+TABLE", re.I | re.S),
    re.compile(r"DROP\s+TEMPORARY\s+TABLE", re.I),
]


async def handle_sql_safety(data: dict) -> dict:
    if data.get("tool_name") != "Bash":
        return {}
    cmd = _get_tool_input(data).get("command", "")
    if "mysql" not in cmd.lower():
        return {}
    for safe in SQL_SAFE_OVERRIDES:
        if safe.search(cmd):
            return {}
    for pattern, label in DANGEROUS_SQL_PATTERNS:
        if pattern.search(cmd):
            return {
                "decision": "block",
                "reason": (
                    f"SQL SAFETY: Blocked dangerous operation: {label}\n"
                    f"Command: {cmd[:200]}\n\n"
                    f"If this is intentional, ask the user to confirm before proceeding."
                ),
            }
    return {}


# ── handle_release_gate_enforce ──────────────────────────────────────────
GATED_PATTERNS = [
    "git push",
    "git tag",
    "gh release create",
    "gh release upload",
    "zip ",
    "7z ",
    "tar ",
    "Compress-Archive",
]
ALWAYS_ALLOWED_GIT = [
    "git push origin master",
    "git push origin main",
]


def _gate_extract_command(command: str) -> str:
    for marker in ("<<'EOF'", '<<"EOF"', "<<EOF", "<<-'EOF'"):
        if marker in command:
            return command[: command.index(marker)]
    return command.split("\n")[0]


def _gate_targets_publishable(command: str) -> bool:
    cmd_lower = command.lower()
    return any(
        ind in cmd_lower
        for ind in ("tools/publishable", "tools\\publishable", "release-gate", "release_gate")
    )


def _gate_is_release_action(command: str) -> bool:
    cmd_line = _gate_extract_command(command).lower().strip()
    for allowed in ALWAYS_ALLOWED_GIT:
        if cmd_line.startswith(allowed):
            return False
    for pattern in GATED_PATTERNS:
        if pattern in cmd_line:
            if pattern == "git push":
                return "--tags" in cmd_line or "refs/tags" in cmd_line
            if pattern in ("zip ", "7z ", "tar ", "Compress-Archive"):
                return _gate_targets_publishable(command)
            return True
    return False


def _gate_read_status() -> "tuple[str, list[str]]":
    if not GATE_STATUS_FILE.exists():
        return ("UNKNOWN", ["No release gate audit has been run. Use /pre-ship first."])
    try:
        d = json.loads(GATE_STATUS_FILE.read_text(encoding="utf-8"))
        return (d.get("status", "UNKNOWN"), d.get("blockers", []))
    except (json.JSONDecodeError, OSError):
        return ("ERROR", ["Could not read release-gate-status.json"])


async def handle_release_gate_enforce(data: dict) -> dict:
    if data.get("tool_name") != "Bash":
        return {}
    cmd = _get_tool_input(data).get("command", "")
    if not cmd or not _gate_is_release_action(cmd):
        return {}
    status, blockers = _gate_read_status()
    if status == "PASS":
        return {}
    msg = f"RELEASE GATE: {status}\n"
    msg += "This action is blocked because the release gate has not passed.\n"
    if blockers:
        msg += "Open blockers:\n"
        for b in blockers[:5]:
            msg += f"  - {b}\n"
        if len(blockers) > 5:
            msg += f"  ... and {len(blockers) - 5} more\n"
    msg += "\nRun /pre-ship to audit, fix blockers, then retry."
    return {"decision": "block", "reason": msg}


# ── handle_sensitive_file_guard ──────────────────────────────────────────
SENSITIVE_PATTERNS = [
    "case_reference/",
    "desktop/excluded/",
    "personal_data_matrix",
    "career_evidence",
    "deep_data_matrix",
    "claude_browser_final_",
]
VOXCORE_ROOT = _normalize_path(str(Path.home() / "VoxCore"))


async def handle_sensitive_file_guard(data: dict) -> dict:
    if data.get("tool_name") not in ("Edit", "Write"):
        return {}
    file_path = _get_tool_input(data).get("file_path", "")
    normalized = _normalize_path(file_path)
    if not normalized.startswith(VOXCORE_ROOT):
        return {}
    for pattern in SENSITIVE_PATTERNS:
        if pattern in normalized:
            return {
                "decision": "block",
                "reason": (
                    f"Blocked: attempting to write sensitive case/personal file inside VoxCore repo. "
                    f"Pattern '{pattern}' matched in path '{file_path}'. "
                    f"These files should stay on Desktop, not in the git-tracked repo."
                ),
            }
    return {}


# ── handle_cpp_build_reminder ────────────────────────────────────────────
async def handle_cpp_build_reminder(data: dict) -> dict:
    global SESSION_CPP_NEW_LINES
    path = _get_tool_input(data).get("file_path", "")
    if not path.endswith((".cpp", ".h")):
        return {}

    messages = []
    messages.append(f"[hook] C++ file modified: {path} — use /build-loop to iterate or build in VS")

    # Feature #1: Set dirty flag for codebase.db auto-refresh
    try:
        CODEBASE_DIRTY_FLAG.parent.mkdir(parents=True, exist_ok=True)
        CODEBASE_DIRTY_FLAG.write_text(
            json.dumps({"dirty_since": datetime.now(timezone.utc).isoformat(), "file": path}),
            encoding="utf-8",
        )
    except Exception:
        log.exception("failed to write codebase dirty flag")

    # Feature #2: Check digest staleness for this file
    basename = Path(path).stem  # e.g., "SpellEffects" from SpellEffects.cpp
    digest_file = DIGEST_CACHE_DIR / f"{basename}.digest.md"
    if digest_file.exists():
        try:
            src_mtime = Path(path).stat().st_mtime if Path(path).exists() else 0
            dig_mtime = digest_file.stat().st_mtime
            if src_mtime > dig_mtime:
                age_hours = (src_mtime - dig_mtime) / 3600
                messages.append(
                    f"[hook] Digest for {basename} is stale (source modified {age_hours:.0f}h after digest). "
                    f"Consider re-generating with: python tools/digest_source.py {basename}"
                )
        except Exception:
            pass

    # Feature #9: Track new C++ lines for Triad reminder
    tool_input = _get_tool_input(data)
    new_str = tool_input.get("new_string", "")
    old_str = tool_input.get("old_string", "")
    if new_str and old_str:
        new_lines = new_str.count("\n") - old_str.count("\n")
        if new_lines > 0:
            SESSION_CPP_NEW_LINES += new_lines
            if SESSION_CPP_NEW_LINES > 100:
                messages.append(
                    f"[hook] {SESSION_CPP_NEW_LINES} new C++ lines this session without a ChatGPT spec. "
                    f"Consider: did the Triad design this? Run `python tools/api_architect/run_architect.py` for a spec."
                )

    return {"systemMessage": "\n".join(messages)}


# ── handle_edit_verifier ─────────────────────────────────────────────────
ADVISORY_ONLY_EXTS = {".md", ".txt", ".json", ".toml", ".yml", ".yaml"}


def _normalize_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return unicodedata.normalize("NFC", text)


def _read_file_for_verify(file_path: str) -> "tuple[str | None, str | None]":
    """Sync helper: returns (content, error) where exactly one is non-None."""
    for encoding in ("utf-8", None, "latin1"):
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read(), None
        except (UnicodeDecodeError, UnicodeError):
            continue
        except (FileNotFoundError, PermissionError, OSError):
            return None, "io_error"
    return None, "decode_error"


async def handle_edit_verifier(data: dict) -> dict:
    if data.get("tool_name") != "Edit":
        return {}
    tool_input = _get_tool_input(data)
    tool_response = data.get("tool_response", {})
    file_path = tool_input.get("file_path", "")
    old_string = tool_input.get("old_string", "")
    new_string = tool_input.get("new_string", "")
    replace_all = tool_input.get("replace_all", False)

    _, ext = os.path.splitext(file_path)
    advisory_only = ext.lower() in ADVISORY_ONLY_EXTS

    if not file_path:
        return {}

    min_chars = int(os.environ.get("EDIT_VERIFY_MIN_CHARS", "3"))
    if new_string and len(new_string.strip()) < min_chars:
        return {}
    if isinstance(tool_response, dict) and not tool_response.get("success", True):
        return {}

    # Blocking file read off the event loop
    content, error = await asyncio.to_thread(_read_file_for_verify, file_path)
    if error == "io_error":
        return {
            "decision": "block",
            "reason": (
                f"Edit verification: Could not read '{file_path}' after the edit. "
                f"The file may not exist or may be locked. "
                f"Please verify the edit applied correctly."
            ),
        }
    if error == "decode_error" or content is None:
        return {
            "decision": "block",
            "reason": (
                f"Edit verification: Could not decode '{file_path}' with any encoding "
                f"(utf-8, system default, latin1). Cannot verify edit."
            ),
        }

    content_norm = _normalize_text(content)
    new_norm = _normalize_text(new_string)
    old_norm = _normalize_text(old_string)

    problems = []
    if new_string and new_norm not in content_norm:
        problems.append(
            f"MISSING NEW CONTENT: The expected new text was not found in "
            f"'{file_path}' after the Edit. The edit may not have applied."
        )
    if old_norm and old_norm in content_norm:
        if replace_all:
            problems.append(
                f"OLD CONTENT STILL PRESENT (replace_all=true): The original text "
                f"still exists in '{file_path}' despite replace_all being set. "
                f"The edit may have failed or matched incorrectly."
            )
        elif old_norm != new_norm and new_norm and new_norm not in content_norm:
            occurrences = content_norm.count(old_norm)
            problems.append(
                f"POSSIBLE EDIT FAILURE: old_string still appears {occurrences} time(s) "
                f"and new_string is missing in '{file_path}'. Read the file to confirm."
            )

    if not problems:
        return {}
    if advisory_only:
        return {}
    return {
        "decision": "block",
        "reason": (
            "Edit verification FAILED:\n"
            + "\n".join(f"  - {p}" for p in problems)
            + "\n\nPlease read the file to verify the edit applied "
            + "correctly before proceeding."
        ),
    }


# ── handle_timestamp_injector ────────────────────────────────────────────
async def handle_timestamp_injector(data: dict) -> dict:
    now = datetime.now()
    timestamp = now.strftime("%Y-%m-%d %H:%M:%S")
    day_of_week = now.strftime("%A")
    return {
        "additionalContext": f"[TIMESTAMP] User message received: {day_of_week} {timestamp}"
    }


# ── handle_prompt_context_injector ───────────────────────────────────────
def _detect_prompt_context(prompt: str) -> "list[str]":
    if len(prompt.strip()) < 10:
        return []
    p = prompt.lower()
    parts: "list[str]" = []
    if any(kw in p for kw in ["transmog", "outfit", "wardrobe", "viewedoutfit", "adt", "idt"]):
        parts.append(
            "CONTEXT: Transmog UI work is ARCHIVED — reimplemented externally. "
            "Historical docs in doc/archive/transmog.md. No active transmog commands or agents."
        )
    if any(kw in p for kw in ["build error", "compile", "linker", "ninja", "cmake", "lnk2"]):
        parts.append(
            "CONTEXT: Build issue detected. Building from Claude Code IS allowed (ninja -j32). "
            "Use /parse-errors to categorize build output. Use /build-loop for iterative build+fix cycles."
        )
    if any(kw in p for kw in ["crash", "server log", "dberror", "assertion", "restart"]):
        parts.append(
            "CONTEXT: Server issue detected. Run /check-logs immediately (it's auto-approved, read-only). "
            "Follow the 4-gate debugging pipeline in .claude/rules/debugging.md"
        )
    if any(
        kw in p
        for kw in [
            "sql",
            "insert into",
            "update ",
            "delete from",
            "alter table",
            "creature_template",
            "hotfixes.",
            "world.",
            "characters.",
        ]
    ):
        parts.append(
            "CONTEXT: SQL work detected. DESCRIBE tables before writing SQL. "
            "Use /new-sql-update for proper filename. Use /apply-sql to apply. "
            "Check doc/session_state.md for multi-tab DB locking."
        )
    if any(kw in p for kw in ["companion", "companionai", "scompanionmgr", ".comp "]):
        parts.append(
            "CONTEXT: Companion system. Files: src/server/game/Companion/, "
            "src/server/scripts/Custom/Companion/. Memory: companion-system.md"
        )
    if any(kw in p for kw in ["spell ", "spellid", "spell_name", "spell_effect", "aura"]):
        parts.append(
            "CONTEXT: Spell work. Custom spells use hotfix tables (spell_name, spell_misc, spell_effect + hotfix_data). "
            "Range 1900003+. Use /lookup-spell to resolve names to IDs. "
            "Don't also add serverside_spell if ID is in spell_name."
        )
    if any(kw in p for kw in ["packet", "opcode", "cmsg", "smsg", "sniff"]):
        parts.append(
            "CONTEXT: Packet analysis. Use /decode-pkt for raw .pkt files, /parse-packet for parsed output. "
            "Delegate to packet-analyzer agent (haiku, read-only) for heavy analysis."
        )
    if any(kw in p for kw in ["creature", "npc ", "cnpc", "customnpc", "spawn"]):
        parts.append(
            "CONTEXT: NPC work. Custom NPCs: CreatureTemplateIdStart=400000. "
            "Use /lookup-creature. creature_template col is 'faction' (not FactionID), "
            "'npcflag' (bigint), spells in creature_template_spell."
        )
    if any(kw in p for kw in [".effect", ".display", "spellvisualkit", "morph", "wmorph"]):
        parts.append(
            "CONTEXT: Visual system. Effects: Noblegarden::EffectsHandler. "
            "Display: RoleplayCore::DisplayHandler. Morph: player_morph_scripts.cpp"
        )

    # Feature #4: FTS5 smart context — query codebase.db for relevant files
    if len(prompt.strip()) > 20 and CODEBASE_DB_PATH.exists():
        fts_context = _fts5_context_lookup(prompt)
        if fts_context:
            parts.append(fts_context)

    return parts


def _fts5_context_lookup(prompt: str) -> "str | None":
    """Query codebase.db FTS5 index for files relevant to the user's prompt.
    Returns a context string with top 3 matching file paths, or None."""
    # Extract meaningful keywords (skip common short words)
    stop_words = {"the", "a", "an", "is", "are", "was", "were", "be", "been", "to", "of",
                  "and", "or", "in", "on", "at", "for", "with", "this", "that", "it", "can",
                  "you", "we", "do", "how", "what", "why", "where", "when", "please", "also",
                  "just", "make", "fix", "add", "get", "set", "use", "run", "check", "look"}
    words = [w for w in re.findall(r"[a-zA-Z_]{3,}", prompt.lower()) if w not in stop_words]
    if len(words) < 1:
        return None
    # Use top 3 keywords for FTS5 query
    query_terms = " ".join(words[:3])
    try:
        conn = sqlite3.connect(str(CODEBASE_DB_PATH), timeout=1.0)
        conn.execute("PRAGMA query_only = ON")
        rows = conn.execute(
            "SELECT name, kind, file_path FROM fts_search WHERE fts_search MATCH ? LIMIT 5",
            (query_terms,),
        ).fetchall()
        conn.close()
        if not rows:
            return None
        # Deduplicate file paths
        seen = set()
        files = []
        for name, kind, file_path in rows:
            if file_path and file_path not in seen:
                seen.add(file_path)
                files.append(f"{file_path} ({kind}: {name})")
        if not files:
            return None
        return "CONTEXT [FTS5]: Relevant codebase files: " + "; ".join(files[:3])
    except Exception:
        return None


async def handle_prompt_context_injector(data: dict) -> dict:
    prompt = data.get("prompt", "")
    if not prompt:
        return {}
    injections = _detect_prompt_context(prompt)
    if not injections:
        return {}
    return {"additionalContext": "\n".join(injections)}


# ── handle_precompact_snapshot ───────────────────────────────────────────
async def handle_precompact_snapshot(data: dict) -> dict:
    """PreCompact: dump recent activity to precompact-state.json so the
    SessionStart compact-reinject hook can restore real context."""
    recent = _read_jsonl_recent(STATS_FILE, minutes=120)
    recent_tools: "list[str]" = []
    recent_files: "list[str]" = []
    for entry in recent:
        tool = entry.get("tool", "")
        if tool:
            recent_tools.append(tool)
        for key in ("file_path", "path", "pattern"):
            if key in entry:
                recent_files.append(entry[key])
                break

    tool_counts = Counter(recent_tools)
    seen: "set[str]" = set()
    unique_files: "list[str]" = []
    for f in reversed(recent_files):
        if f not in seen:
            seen.add(f)
            unique_files.append(f)
    unique_files = list(reversed(unique_files[-20:]))

    cpp_files = [f for f in unique_files if f.endswith((".cpp", ".h"))]
    sql_files = [f for f in unique_files if f.endswith(".sql")]
    transmog_files = [
        f for f in unique_files if "transmog" in f.lower() or "display" in f.lower()
    ]

    work_signals: "list[str]" = []
    if cpp_files:
        work_signals.append(
            f"C++ editing: {', '.join(os.path.basename(f) for f in cpp_files[:5])}"
        )
    if sql_files:
        work_signals.append(
            f"SQL work: {', '.join(os.path.basename(f) for f in sql_files[:5])}"
        )
    if transmog_files:
        work_signals.append("Transmog system work detected")
    if tool_counts.get("Agent", 0) > 0:
        work_signals.append(f"Spawned {tool_counts['Agent']} subagent(s)")

    snapshot = {
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "trigger": data.get("trigger", "unknown"),
        "recent_files": unique_files,
        "tool_usage": dict(tool_counts.most_common(10)),
        "work_signals": work_signals,
        "cpp_files_touched": cpp_files,
        "sql_files_touched": sql_files,
    }
    try:
        PRECOMPACT_STATE.parent.mkdir(parents=True, exist_ok=True)
        PRECOMPACT_STATE.write_text(json.dumps(snapshot, indent=2), encoding="utf-8")
    except Exception:
        log.exception("precompact_snapshot write failed")

    if work_signals:
        return {
            "systemMessage": f"Pre-compaction snapshot saved: {', '.join(work_signals)}"
        }
    return {}


# ── handle_stop_verify ───────────────────────────────────────────────────
async def handle_stop_verify(data: dict) -> dict:
    """Stop hook: VoxCore-specific workflow enforcement. Advisory only."""
    transcript_lower = str(data.get("transcript_suffix", "")).lower()
    recent = _read_jsonl_recent(STATS_FILE, minutes=120)

    cpp_edits: "list[str]" = []
    sql_edits: "list[str]" = []
    shared_edits: "list[str]" = []
    for entry in recent:
        file_path = entry.get("file_path", "") or entry.get("path", "")
        if not file_path:
            continue
        lower = file_path.lower()
        if lower.endswith((".cpp", ".h")):
            cpp_edits.append(file_path)
        elif lower.endswith(".sql"):
            sql_edits.append(file_path)
        if any(
            shared in lower
            for shared in (
                "session_state.md",
                "custom_script_loader",
                "roleplay.h",
                "roleplay.cpp",
            )
        ):
            shared_edits.append(file_path)

    reminders: "list[str]" = []
    if cpp_edits:
        names = list({os.path.basename(f) for f in cpp_edits})[:5]
        reminders.append(
            f"C++ files edited ({', '.join(names)}) — remind user to build in Visual Studio."
        )
    if sql_edits:
        names = list({os.path.basename(f) for f in sql_edits})[:5]
        reminders.append(
            f"SQL files touched ({', '.join(names)}) — remind about /apply-sql or pending SQL pipeline."
        )
    wrap_indicators = (
        "that should do it",
        "all done",
        "everything is complete",
        "let me know if",
        "anything else",
        "good to go",
        "that covers",
        "we're done",
        "wrapping up",
    )
    if any(ind in transcript_lower for ind in wrap_indicators):
        reminders.append(
            "Session appears to be ending — run /wrap-up to commit, push, sync bridge, and update memory."
        )
    if shared_edits:
        names = list({os.path.basename(f) for f in shared_edits})
        reminders.append(
            f"Shared files modified ({', '.join(names)}) — check doc/session_state.md for multi-tab coordination."
        )

    if not reminders:
        return {}
    return {
        "systemMessage": "WORKFLOW CHECK:\n" + "\n".join(f"  - {r}" for r in reminders)
    }


# ── handle_docx_auto_extract ─────────────────────────────────────────────
async def handle_docx_auto_extract(data: dict) -> dict:
    """PostToolUseFailure(Read) on .docx — extract text via python-docx."""
    path = _get_tool_input(data).get("file_path", "")
    if not path.lower().endswith((".docx", ".doc")):
        return {}
    if not os.path.isfile(path):
        return {}
    if path.lower().endswith(".doc") and not path.lower().endswith(".docx"):
        return {
            "systemMessage": (
                f"[hook] {os.path.basename(path)} is old .doc format — python-docx can't read it. "
                f"Open in Word or convert to .docx first."
            )
        }
    try:
        from docx import Document  # type: ignore  # lazy import
    except Exception:
        return {"systemMessage": "[hook] python-docx not installed; cannot auto-extract .docx"}
    try:
        cache_dir = Path.home() / "VoxCore" / ".claude" / "cache" / "docx_extractions"
        cache_dir.mkdir(parents=True, exist_ok=True)
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + "\t".join(cell.text for cell in row.cells)
        basename = os.path.splitext(os.path.basename(path))[0]
        out_path = cache_dir / f"{basename}.txt"
        out_path.write_text(text, encoding="utf-8")
        return {
            "systemMessage": (
                f"[hook] Auto-extracted .docx to readable text: {out_path}\n"
                f"[hook] Use Read on that .txt path to see the contents."
            )
        }
    except Exception as e:
        log.exception("docx extract failed")
        return {"systemMessage": f"[hook] docx extraction failed: {e}"}


# ═══════════════════════════════════════════════════════════════════════
# SECTION F — ASYNC ADVISORY HANDLERS
# These return {} immediately and run real work in asyncio.create_task()
# ═══════════════════════════════════════════════════════════════════════

# ── handle_session_stats ─────────────────────────────────────────────────
async def handle_session_stats(data: dict) -> dict:
    asyncio.create_task(_session_stats_write(data))
    return {}


async def _session_stats_write(data: dict) -> None:
    entry = {
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "event": data.get("hook_event_name", "unknown"),
        "tool": data.get("tool_name", ""),
        "session": data.get("session_id", ""),
    }
    tool_input = _get_tool_input(data)
    for key in ("file_path", "path", "pattern"):
        if key in tool_input:
            entry[key] = tool_input[key]
            break
    try:
        STATS_FILE.parent.mkdir(parents=True, exist_ok=True)
        with open(STATS_FILE, "a", encoding="utf-8") as f:
            f.write(json.dumps(entry) + "\n")
    except Exception:
        log.exception("session_stats write failed")


# ── handle_release_gate_revalidate ───────────────────────────────────────
async def handle_release_gate_revalidate(data: dict) -> dict:
    """PostToolUse(Write|Edit): mark gate STALE if a publishable file changed."""
    if data.get("tool_name") not in ("Write", "Edit"):
        return {}
    file_path = _get_tool_input(data).get("file_path", "")
    if not file_path:
        return {}
    normalized = file_path.replace("\\", "/")
    if "tools/publishable/" not in normalized:
        return {}
    msg = ""
    if GATE_STATUS_FILE.exists():
        try:
            d = json.loads(GATE_STATUS_FILE.read_text(encoding="utf-8"))
            if d.get("status") == "PASS":
                d["status"] = "STALE"
                d["stale_reason"] = f"File edited after gate pass: {file_path}"
                GATE_STATUS_FILE.write_text(
                    json.dumps(d, indent=2, ensure_ascii=False), encoding="utf-8"
                )
                msg = (
                    "Release gate: File in publishable/ was edited. "
                    "Gate status is now STALE. Re-run /pre-ship before shipping."
                )
        except (json.JSONDecodeError, OSError):
            log.exception("release_gate_revalidate read/write failed")
    if msg:
        return {"systemMessage": msg}
    return {}


# ── handle_large_file_guard ──────────────────────────────────────────────
LARGE_FILE_THRESHOLD = 3000
LARGE_FILE_KNOWN_OK = {"CLAUDE.md", "worldserver.conf.dist", "worldserver.conf"}


def _count_lines(file_path: str) -> int:
    """Sync line count helper. Returns -1 on error."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return sum(1 for _ in f)
    except (FileNotFoundError, PermissionError, OSError):
        return -1


async def handle_large_file_guard(data: dict) -> dict:
    if data.get("tool_name") != "Read":
        return {}
    tool_input = _get_tool_input(data)
    file_path = tool_input.get("file_path", "")
    if not file_path:
        return {}
    basename = os.path.basename(file_path)
    if basename in LARGE_FILE_KNOWN_OK:
        return {}
    if tool_input.get("offset") or tool_input.get("limit"):
        return {}
    line_count = await asyncio.to_thread(_count_lines, file_path)
    if line_count < 0 or line_count <= LARGE_FILE_THRESHOLD:
        return {}
    return {
        "systemMessage": (
            f"[hook] Large file read: {basename} has {line_count} lines. "
            f"Consider using offset/limit parameters to read specific sections "
            f"and save context tokens."
        )
    }


# ── handle_sync_on_git ───────────────────────────────────────────────────
async def handle_sync_on_git(data: dict) -> dict:
    cmd = _get_tool_input(data).get("command", "")
    if any(kw in cmd for kw in ("git commit", "git push", "git merge")):
        _spawn_detached([sys.executable, "C:/Users/atayl/cowork/sync_bridge.py"])
    return {}


# ── handle_subagent_complete ─────────────────────────────────────────────
async def handle_subagent_complete(data: dict) -> dict:
    asyncio.create_task(_subagent_complete_work(data))
    return {}


async def _subagent_complete_work(data: dict) -> None:
    entry = {
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "event": "SubagentStop",
        "session": data.get("session_id", ""),
    }
    try:
        STATS_FILE.parent.mkdir(parents=True, exist_ok=True)
        with open(STATS_FILE, "a", encoding="utf-8") as f:
            f.write(json.dumps(entry) + "\n")
    except Exception:
        log.exception("subagent_complete stats write failed")
    burnttoast = (
        "try { New-BurntToastNotification "
        "-Text 'Subagent Complete', 'Background agent finished' "
        "-ExpirationTime ([datetime]::Now.AddSeconds(6)) "
        "-Silent "
        "} catch { }"
    )
    _spawn_detached(["powershell.exe", "-NoProfile", "-Command", burnttoast])


# ── handle_notification_toast ────────────────────────────────────────────
async def handle_notification_toast(data: dict) -> dict:
    _spawn_notification_toast(data)
    return {}


def _spawn_notification_toast(data: dict) -> None:
    hook_event = data.get("hook_event_name", "Notification")
    tool_name = data.get("tool_name", "")
    msg_raw = data.get("message", "")

    if "permission" in hook_event.lower() or "permission" in str(data).lower():
        title = "Permission Needed"
        message = f"{tool_name}" if tool_name else "Action requires approval"
        duration = 30
        sound = True
    elif hook_event == "Stop":
        title = "Task Complete"
        message = "Claude Code finished working"
        duration = 8
        sound = False
    elif hook_event == "PostToolUseFailure":
        title = "Tool Failed"
        message = (
            f"{tool_name} encountered an error" if tool_name else "A tool encountered an error"
        )
        duration = 15
        sound = True
    elif "idle" in str(data).lower() or "waiting" in msg_raw.lower():
        title = "Waiting for Input"
        message = "Claude Code is waiting for your response"
        duration = 20
        sound = True
    else:
        title = "Claude Code"
        message = msg_raw if msg_raw else "Needs your attention"
        duration = 10
        sound = False

    title_ps = title.replace('"', '`"').replace("'", "''")
    message_ps = message.replace('"', '`"').replace("'", "''")
    sound_param = "" if sound else " -Silent"
    burnttoast_cmd = (
        f"New-BurntToastNotification "
        f"-Text '{title_ps}', '{message_ps}' "
        f"-ExpirationTime ([datetime]::Now.AddSeconds({duration}))"
        f"{sound_param}"
    )
    icon = "Warning" if sound else "Info"
    fallback_cmd = (
        f"Add-Type -AssemblyName System.Windows.Forms; "
        f"$n = New-Object System.Windows.Forms.NotifyIcon; "
        f"$n.Icon = [System.Drawing.SystemIcons]::{icon}; "
        f"$n.Visible = $true; "
        f"$n.ShowBalloonTip(5000, '{title_ps}', '{message_ps}', "
        f"[System.Windows.Forms.ToolTipIcon]::{icon}); "
        f"Start-Sleep -Seconds 6; $n.Dispose()"
    )
    ps_cmd = f"try {{ {burnttoast_cmd} }} catch {{ {fallback_cmd} }}"
    _spawn_detached(["powershell.exe", "-NoProfile", "-Command", ps_cmd])


# ── handle_file_changed_monitor ──────────────────────────────────────────
SESSION_FILES = ("session_state.md", "Central_Brain.md", "todo.md")


async def handle_file_changed_monitor(data: dict) -> dict:
    file_path = data.get("file_path", "")
    change_type = data.get("change_type", "modified")
    if not file_path:
        return {}
    skip_patterns = (".git/", "node_modules/", "__pycache__/", ".tmp", "~")
    if any(p in file_path for p in skip_patterns):
        return {}
    name = os.path.basename(file_path)
    ext = os.path.splitext(file_path)[1].lower()
    if name in SESSION_FILES:
        category = "coordination"
    elif ext in (".sql",):
        category = "sql"
    elif ext in (".cpp", ".h", ".hpp"):
        category = "cpp"
    elif "publishable" in file_path:
        category = "addon"
    else:
        category = "other"
    try:
        rel_path = os.path.relpath(file_path, str(PROJECT_DIR))
    except ValueError:
        rel_path = file_path

    # Log to session stats (fire-and-forget)
    asyncio.create_task(
        _file_changed_log(rel_path, change_type, category, data.get("session_id", ""))
    )

    if category == "coordination":
        return {
            "additionalContext": (
                f"[EXTERNAL CHANGE] {rel_path} was {change_type} by another process. "
                f"Re-read it before making decisions based on its content."
            )
        }
    if category == "sql" and "pending" in file_path.lower():
        return {
            "additionalContext": (
                f"[EXTERNAL CHANGE] SQL file {rel_path} was {change_type} externally. "
                f"Another tab may have applied or created SQL."
            )
        }
    if category == "cpp":
        return {
            "additionalContext": (
                f"[EXTERNAL CHANGE] C++ source {rel_path} was {change_type} externally. "
                f"A rebuild may be needed if you were planning to build."
            )
        }
    return {}


async def _file_changed_log(rel_path: str, change_type: str, category: str, session_id: str) -> None:
    entry = {
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "event": "FileChanged",
        "change_type": change_type,
        "file_path": rel_path,
        "category": category,
        "session": session_id,
    }
    try:
        STATS_FILE.parent.mkdir(parents=True, exist_ok=True)
        with open(STATS_FILE, "a", encoding="utf-8") as f:
            f.write(json.dumps(entry) + "\n")
    except Exception:
        log.exception("file_changed log failed")


# ── handle_cowork_sync ───────────────────────────────────────────────────
COWORK_SYNC_SCRIPT = "C:/Users/atayl/cowork/sync_bridge.py"


async def handle_cowork_sync(data: dict) -> dict:
    asyncio.create_task(_cowork_sync_task())
    return {}


async def _cowork_sync_task() -> None:
    try:
        proc = await asyncio.create_subprocess_exec(
            sys.executable,
            COWORK_SYNC_SCRIPT,
            "--full",
            stdout=asyncio.subprocess.DEVNULL,
            stderr=asyncio.subprocess.DEVNULL,
        )
        try:
            await asyncio.wait_for(proc.wait(), timeout=60.0)
        except asyncio.TimeoutError:
            log.warning("cowork_sync timed out after 60s — killing")
            try:
                proc.kill()
            except Exception:
                pass
    except FileNotFoundError:
        log.warning("cowork_sync: %s not found", COWORK_SYNC_SCRIPT)
    except Exception:
        log.exception("cowork_sync error")


# ── handle_codebase_refresh ─────────────────────────────────────────────

async def handle_codebase_refresh(data: dict) -> dict:
    """Called by durable cron job. Checks dirty flag, runs incremental rebuild if needed."""
    if not CODEBASE_DIRTY_FLAG.exists():
        return {"systemMessage": "[cron] Codebase index is clean — no rebuild needed."}
    try:
        flag_data = json.loads(CODEBASE_DIRTY_FLAG.read_text(encoding="utf-8"))
        dirty_since = flag_data.get("dirty_since", "unknown")
    except Exception:
        dirty_since = "unknown"

    build_script = PROJECT_DIR / "tools" / "build_code_index.py"
    if not build_script.exists():
        return {"systemMessage": "[cron] build_code_index.py not found — skipping rebuild."}

    # Run incremental rebuild in background
    asyncio.create_task(_run_codebase_rebuild(build_script, dirty_since))
    return {"systemMessage": f"[cron] Codebase dirty since {dirty_since} — incremental rebuild started."}


async def _run_codebase_rebuild(script: Path, dirty_since: str) -> None:
    try:
        proc = await asyncio.create_subprocess_exec(
            sys.executable, str(script), "--layers", "1-6",
            cwd=str(PROJECT_DIR),
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        try:
            stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=120.0)
            # Clear dirty flag on success
            if proc.returncode == 0:
                try:
                    CODEBASE_DIRTY_FLAG.unlink()
                except Exception:
                    pass
                log.info("codebase_refresh: rebuild complete (dirty since %s)", dirty_since)
            else:
                log.warning("codebase_refresh: rebuild exited %d", proc.returncode)
        except asyncio.TimeoutError:
            log.warning("codebase_refresh: rebuild timed out after 120s")
            try:
                proc.kill()
            except Exception:
                pass
    except Exception:
        log.exception("codebase_refresh: rebuild error")


# ═══════════════════════════════════════════════════════════════════════
# SECTION G — Route table
# ═══════════════════════════════════════════════════════════════════════

async def handle_block_recurring_cron(data: dict) -> dict:
    """Block recurring CronCreate calls — they freeze tabs.

    One-shot reminders (recurring: false) are allowed.
    Recurring jobs (recurring: true, the default) are blocked.
    """
    tool_input = data.get("tool_input", {})

    # CronCreate defaults recurring to True if not specified
    is_recurring = tool_input.get("recurring", True)

    if is_recurring:
        log.warning("BLOCKED recurring CronCreate — freezes tabs")
        return {
            "decision": "block",
            "reason": (
                "Recurring cron jobs are BANNED — they freeze Claude Code tabs by "
                "firing prompts into the REPL when idle. Use `recurring: false` for "
                "one-shot reminders only. For recurring work, use Windows Task "
                "Scheduler + scripts that write results to files."
            ),
        }

    # One-shot is fine
    log.info("Allowed one-shot CronCreate")
    return {}


ROUTE_TABLE = {
    "/hook/block-recurring-cron": handle_block_recurring_cron,
    "/hook/sql-safety": handle_sql_safety,
    "/hook/release-gate-enforce": handle_release_gate_enforce,
    "/hook/sensitive-file-guard": handle_sensitive_file_guard,
    "/hook/cpp-build-reminder": handle_cpp_build_reminder,
    "/hook/edit-verifier": handle_edit_verifier,
    "/hook/release-gate-revalidate": handle_release_gate_revalidate,
    "/hook/large-file-guard": handle_large_file_guard,
    "/hook/sync-on-git": handle_sync_on_git,
    "/hook/session-stats": handle_session_stats,
    "/hook/timestamp-injector": handle_timestamp_injector,
    "/hook/prompt-context-injector": handle_prompt_context_injector,
    "/hook/precompact-snapshot": handle_precompact_snapshot,
    "/hook/subagent-complete": handle_subagent_complete,
    "/hook/notification-toast": handle_notification_toast,
    "/hook/stop-verify": handle_stop_verify,
    "/hook/docx-auto-extract": handle_docx_auto_extract,
    "/hook/file-changed-monitor": handle_file_changed_monitor,
    "/hook/cowork-sync": handle_cowork_sync,
    "/hook/codebase-refresh": handle_codebase_refresh,
}


# ═══════════════════════════════════════════════════════════════════════
# SECTION H — Lifecycle (PID file, daemon liveness, shutdown)
# ═══════════════════════════════════════════════════════════════════════

PROCESS_QUERY_LIMITED_INFORMATION = 0x1000


def _is_pid_alive(pid: int) -> bool:
    try:
        handle = ctypes.windll.kernel32.OpenProcess(
            PROCESS_QUERY_LIMITED_INFORMATION, False, pid
        )
        if handle:
            ctypes.windll.kernel32.CloseHandle(handle)
            return True
    except Exception:
        pass
    return False


def is_daemon_already_running() -> bool:
    if not PID_FILE.exists():
        return False
    try:
        pid = int(PID_FILE.read_text(encoding="ascii").strip())
    except (ValueError, OSError):
        return False
    if pid == os.getpid():
        return False  # our own PID file from a previous boot — not "another" instance
    if not _is_pid_alive(pid):
        return False
    # PID is alive — but is it actually listening? (zombie check)
    try:
        conn = http.client.HTTPConnection(HOST, PORT, timeout=0.5)
        conn.request("GET", "/health")
        r = conn.getresponse()
        if r.status == 200:
            return True
    except Exception:
        pass
    # PID alive but not listening — stale zombie. Clean up.
    log.warning("PID %d alive but not listening on port %d — removing stale PID file", pid, PORT)
    try:
        PID_FILE.unlink()
    except Exception:
        pass
    return False


def write_pid_file() -> None:
    PID_FILE.parent.mkdir(parents=True, exist_ok=True)
    PID_FILE.write_text(str(os.getpid()), encoding="ascii")


def delete_pid_file() -> None:
    try:
        PID_FILE.unlink()
    except FileNotFoundError:
        pass


# ═══════════════════════════════════════════════════════════════════════
# SECTION I — main
# ═══════════════════════════════════════════════════════════════════════

async def main() -> None:
    global SHUTDOWN_EVENT
    SHUTDOWN_EVENT = asyncio.Event()
    loop = asyncio.get_running_loop()

    def _handle_signal(signum, frame):
        # signal handlers run on the main thread; schedule the event set onto the loop
        loop.call_soon_threadsafe(SHUTDOWN_EVENT.set)

    signal.signal(signal.SIGINT, _handle_signal)
    signal.signal(signal.SIGTERM, _handle_signal)

    if is_daemon_already_running():
        log.warning("Daemon already running per PID file; exiting.")
        return

    try:
        server = await asyncio.start_server(client_connected, HOST, PORT)
    except OSError as e:
        log.error("Cannot bind %s:%d — %s", HOST, PORT, e)
        return

    write_pid_file()
    log.info(
        "Hook daemon v%s started on %s:%d, PID %d, %d routes",
        DAEMON_VERSION,
        HOST,
        PORT,
        os.getpid(),
        len(ROUTE_TABLE),
    )

    async with server:
        await SHUTDOWN_EVENT.wait()

    if IN_FLIGHT:
        log.info("Draining %d in-flight requests...", len(IN_FLIGHT))
        try:
            await asyncio.wait(IN_FLIGHT, timeout=5.0)
        except Exception:
            log.exception("drain wait error")

    delete_pid_file()
    log.info("Hook daemon stopped cleanly")


def _cli_health() -> int:
    """Print daemon health to stdout. Exit 0 if healthy, 1 otherwise."""
    try:
        conn = http.client.HTTPConnection(HOST, PORT, timeout=0.5)
        conn.request("GET", "/health")
        r = conn.getresponse()
        body = r.read().decode("utf-8")
        print(body)
        return 0 if r.status == 200 else 1
    except Exception as e:
        print(json.dumps({"status": "unreachable", "error": str(e)}))
        return 1


def _cli_ensure() -> int:
    """Start daemon if not already running. Used as bootstrap from command-type hooks.
    Returns 0 immediately — either daemon was already running or we spawned it detached."""
    if is_daemon_already_running():
        return 0
    # Also check port reachability (PID file may be stale)
    try:
        conn = http.client.HTTPConnection(HOST, PORT, timeout=0.3)
        conn.request("GET", "/health")
        r = conn.getresponse()
        if r.status == 200:
            return 0  # Running but PID file missing — that's fine
    except Exception:
        pass  # Not running — proceed to start
    # Spawn self detached
    DETACHED_PROCESS = 0x00000008
    CREATE_NEW_PROCESS_GROUP = 0x00000200
    try:
        subprocess.Popen(
            [sys.executable, str(Path(__file__).resolve())],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            creationflags=DETACHED_PROCESS | CREATE_NEW_PROCESS_GROUP,
            close_fds=True,
        )
    except Exception as e:
        print(f"Failed to start hook daemon: {e}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    if "--health" in sys.argv:
        sys.exit(_cli_health())
    if "--version" in sys.argv:
        print(DAEMON_VERSION)
        sys.exit(0)
    if "--ensure" in sys.argv:
        sys.exit(_cli_ensure())
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        pass
